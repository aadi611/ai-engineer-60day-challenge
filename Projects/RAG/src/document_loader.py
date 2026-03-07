"""
Document Loader Module
Handles loading and preprocessing of various document formats
"""

import os
from pathlib import Path
from typing import List, Dict, Optional
from loguru import logger
import mimetypes

# Document processing imports
import pypdf
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup


class Document:
    """Represents a processed document with metadata"""
    
    def __init__(
        self,
        content: str,
        metadata: Optional[Dict] = None,
        doc_id: Optional[str] = None
    ):
        self.content = content
        self.metadata = metadata or {}
        self.doc_id = doc_id or self._generate_id()
        
    def _generate_id(self) -> str:
        """Generate unique document ID"""
        import hashlib
        return hashlib.md5(self.content.encode()).hexdigest()
    
    def __repr__(self) -> str:
        return f"Document(id={self.doc_id}, length={len(self.content)})"


class DocumentLoader:
    """
    Production-grade document loader supporting multiple formats
    """
    
    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    }
    
    def __init__(self, max_file_size_mb: int = 50):
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        logger.info(f"DocumentLoader initialized with max file size: {max_file_size_mb}MB")
    
    def load(self, file_path: str) -> Document:
        """
        Load a single document from file path
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document object with content and metadata
            
        Raises:
            ValueError: If file format not supported or file too large
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        # Validate file exists
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Validate file size
        file_size = path.stat().st_size
        if file_size > self.max_file_size_bytes:
            raise ValueError(
                f"File too large: {file_size / 1024 / 1024:.2f}MB "
                f"(max: {self.max_file_size_bytes / 1024 / 1024:.2f}MB)"
            )
        
        # Validate file format
        extension = path.suffix.lower()
        if extension not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported: {list(self.SUPPORTED_FORMATS.keys())}"
            )
        
        logger.info(f"Loading document: {file_path}")
        
        # Load based on file type
        if extension == '.pdf':
            content = self._load_pdf(path)
        elif extension == '.txt':
            content = self._load_txt(path)
        elif extension == '.md':
            content = self._load_markdown(path)
        elif extension == '.docx':
            content = self._load_docx(path)
        else:
            raise ValueError(f"Handler not implemented for: {extension}")
        
        # Create metadata
        metadata = {
            'source': str(path),
            'file_name': path.name,
            'file_type': extension,
            'file_size': file_size,
        }
        
        document = Document(content=content, metadata=metadata)
        # Set doc_id in metadata so chunker can use it for unique chunk IDs
        document.metadata['doc_id'] = document.doc_id
        logger.success(f"Loaded document: {path.name} ({len(content)} characters)")
        
        return document
    
    def load_directory(
        self,
        directory_path: str,
        recursive: bool = True
    ) -> List[Document]:
        """
        Load all supported documents from a directory
        
        Args:
            directory_path: Path to directory
            recursive: Whether to search subdirectories
            
        Returns:
            List of Document objects
        """
        path = Path(directory_path)
        
        if not path.is_dir():
            raise ValueError(f"Not a directory: {directory_path}")
        
        logger.info(f"Loading documents from: {directory_path}")
        
        documents = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    doc = self.load(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")
                    continue
        
        logger.success(f"Loaded {len(documents)} documents from directory")
        return documents
    
    def _load_pdf(self, path: Path) -> str:
        """Load PDF document"""
        try:
            with open(path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                text = []
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
                return "\n\n".join(text)
        except Exception as e:
            logger.error(f"Error loading PDF {path}: {e}")
            raise
    
    def _load_txt(self, path: Path) -> str:
        """Load text document"""
        try:
            with open(path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _load_markdown(self, path: Path) -> str:
        """Load markdown document and convert to text"""
        try:
            with open(path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to HTML then to plain text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
        except Exception as e:
            logger.error(f"Error loading Markdown {path}: {e}")
            raise
    
    def _load_docx(self, path: Path) -> str:
        """Load DOCX document"""
        try:
            doc = DocxDocument(path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return "\n\n".join(text)
        except Exception as e:
            logger.error(f"Error loading DOCX {path}: {e}")
            raise


def preprocess_text(text: str) -> str:
    """
    Preprocess text for better chunking and retrieval
    
    Args:
        text: Raw text content
        
    Returns:
        Preprocessed text
    """
    # Remove excessive whitespace
    text = ' '.join(text.split())
    
    # Normalize line breaks
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Remove multiple consecutive newlines
    while '\n\n\n' in text:
        text = text.replace('\n\n\n', '\n\n')
    
    return text.strip()


if __name__ == "__main__":
    # Test the document loader
    loader = DocumentLoader()
    
    # Example usage
    print("DocumentLoader module ready")
    print(f"Supported formats: {list(DocumentLoader.SUPPORTED_FORMATS.keys())}")
